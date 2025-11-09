import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def _p(doc, text="", bold=False, size=12, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if align:
        p.alignment = align
    return p


def add_logo_to_header(section, logo_bytes: bytes):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(logo_bytes), width=Inches(1.0), height=Inches(1.0))


def _force_cell_vertical_center(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(r'<w:vAlign %s w:val="center"/>' % nsdecls('w')))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_materiaalstaat_page(doc: Document, materialen: list[dict]):
    """Voegt materiaalstaat toe op eigen pagina."""
    doc.add_page_break()

    title_p = doc.add_paragraph()
    run = title_p.add_run("Materiaalstaat")
    run.font.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    _p(doc, "")

    cols = ["Nummer", "Aantal", "Benaming", "Lengte", "Breedte", "Dikte", "Materiaal"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    # header
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        cell = hdr_cells[i]
        cell.text = col_name
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(12)
        cell._element.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w")))
        )
        _force_cell_vertical_center(cell)

    # data
    for item in materialen:
        row_cells = table.add_row().cells
        for j, key in enumerate(cols):
            value = item.get(key, "")
            cell = row_cells[j]
            cell.text = value
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(12)
            _force_cell_vertical_center(cell)

        # rijhoogte vergroten
        tr = row_cells[0]._tc.getparent()
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(r'<w:trHeight {} w:val="600"/>'.format(nsdecls("w")))
        trPr.append(trHeight)

    _p(doc, "")
    _p(doc, "")


def add_cover_page(
    doc: Document,
    *,
    opdracht_titel: str,
    vak: str,
    profieldeel: str,
    docent: str,
    duur: str,
    logo: bytes = None,
    cover_bytes: bytes = None,
):
    if logo:
        add_logo_to_header(doc.sections[0], logo)

    _p(doc, "Opdracht :", bold=True, size=14)
    _p(doc, opdracht_titel or " ", bold=True, size=28)
    _p(doc, "")
    _p(doc, vak or "", bold=True, size=14)

    # profieldeel
    if profieldeel:
        _p(doc, f"Keuze/profieldeel: {profieldeel}", size=12)
    else:
        _p(doc, "Keuze/profieldeel:", size=12)

    # docent
    if docent:
        _p(doc, f"Docent: {docent}", size=12)
    else:
        _p(doc, "Docent:", size=12)

    # duur
    if duur:
        _p(doc, f"Duur van de opdracht:     {duur}", size=12)
    else:
        _p(doc, "Duur van de opdracht:", size=12)

    _p(doc, "")

    # cover-afbeelding
    if cover_bytes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(io.BytesIO(cover_bytes), width=Inches(4.5))
        _p(doc, "")

    # naam / klas
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Naam:"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "Klas:"
    table.rows[1].cells[1].text = ""

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(12)

    _p(doc, "")
    _p(doc, "")


def build_workbook_docx_front_and_steps(meta: dict, steps: list[dict]) -> io.BytesIO:
    """
    - Voorpagina
    - (optioneel) Materiaalstaat
    - Elke stap/pagina op EIGEN pagina
    """
    doc = Document()

    add_cover_page(
        doc,
        opdracht_titel=meta.get("opdracht_titel", ""),
        vak=meta.get("vak", "BWI"),
        profieldeel=meta.get("profieldeel", ""),
        docent=meta.get("docent", ""),
        duur=meta.get("duur", ""),
        logo=meta.get("logo"),
        cover_bytes=meta.get("cover_bytes"),
    )

    if meta.get("include_materiaalstaat"):
        add_materiaalstaat_page(doc, meta.get("materialen", []))

    # nu elke “stap” / pagina op z’n eigen pagina
    for idx, step in enumerate(steps):
        # altijd page break vóór de pagina (behalve als er helemaal geen materiaalstaat was en dit de eerste is?)
        doc.add_page_break()

        # titel
        if step.get("title"):
            doc.add_heading(step["title"], level=1)

        # tekstblokken
        for txt in step.get("text_blocks", []):
            _p(doc, txt, size=11)

        # afbeeldingen
        for img_bytes in step.get("images", []):
            if img_bytes:
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(4.5))
                _p(doc, "")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out



