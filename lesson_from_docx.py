import os
import io
import json
import re
from docx import Document
from openai import OpenAI, RateLimitError, APIError


def docx_to_blocks(file_like):
    """
    Leest het .docx bestand en maakt blokken: [{"title": ..., "body": ...}, ...]
    Kop = heading / vet / ALL CAPS
    """
    doc = Document(file_like)
    blocks = []
    current_title = None
    current_body = []

    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue

        is_heading = (
            (para.style and para.style.name and para.style.name.lower().startswith("heading"))
            or any(r.bold for r in para.runs)
            or (len(txt) <= 50 and txt.upper() == txt)
        )

        if is_heading:
            if current_title or current_body:
                blocks.append({
                    "title": current_title,
                    "body": "\n".join(current_body).strip()
                })
            current_title = txt
            current_body = []
        else:
            current_body.append(txt)

    if current_title or current_body:
        blocks.append({
            "title": current_title,
            "body": "\n".join(current_body).strip()
        })

    # als er écht niks in staat, stoppen we meteen
    if not blocks:
        raise RuntimeError("Het Word-bestand bevatte geen herkenbare tekst/onderdelen.")
    return blocks


def ai_blocks_to_slides(blocks):
    """
    1 AI-call: we geven alle blokken, AI geeft slides terug.
    GEEN fallback: als dit faalt → raise.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY ontbreekt. Zet je sleutel in de omgeving.")

    client = OpenAI(api_key=api_key)

    # blokken in prompt
    parts = []
    for i, b in enumerate(blocks, start=1):
        parts.append(
            f"### Onderdeel {i}\nKop: {b.get('title') or ''}\nTekst:\n{b.get('body') or ''}\n"
        )
    joined = "\n\n".join(parts)

    prompt = f"""
Je krijgt hieronder meerdere onderdelen uit een les over installatietechniek / sanitair.
Maak hier lesonderdelen van voor een VMBO-klas (basis/kader/GL).

Voor ELK onderdeel:
- maak 1 korte, begrijpelijke titel (max 8 woorden)
- maak 2 of 3 korte, vertellende zinnen in de je-vorm
- maak 1 controlevraag
- gebruik eenvoudige woorden
- herhaal de titel NIET in de tekst

Geef ALLEEN geldig JSON terug in dit formaat:

{{
  "slides": [
    {{
      "title": "…",
      "text": ["…", "…", "…"],
      "check": "…"
    }}
  ]
}}

Hier zijn de onderdelen:
{joined}
"""

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
    except RateLimitError as e:
        # hier géén fallback, gewoon stoppen
        raise RuntimeError("AI-limiet bereikt bij OpenAI. Probeer het later opnieuw.") from e
    except APIError as e:
        # bv insufficient_quota
        raise RuntimeError(f"AI gaf een fout terug: {e}") from e

    try:
        data = json.loads(resp.choices[0].message.content)
    except Exception as e:
        raise RuntimeError("AI gaf geen geldig JSON terug.") from e

    slides = data.get("slides")
    if not slides:
        raise RuntimeError("AI gaf geen slides terug.")
    return slides


def build_word_from_slides(slides):
    """
    Maakt een .docx met dit format:

    LessonUp-les (gegenereerd)
    1️⃣ Titel
    Uitleg
    ...
    Vraag
    ...
    (lege regel)
    """
    doc = Document()
    doc.add_heading("LessonUp-les (gegenereerd met AI)", level=0)

    for idx, slide in enumerate(slides, start=1):
        title = slide.get("title") or f"Onderdeel {idx}"
        text_lines = slide.get("text") or []
        check = slide.get("check") or ""

        # nummer + titel
        doc.add_heading(f"{idx}️⃣ {title}", level=1)

        # Uitleg
        p = doc.add_paragraph()
        p.add_run("Uitleg").bold = True

        for line in text_lines:
            doc.add_paragraph(line)

        if check:
            q = doc.add_paragraph()
            q.add_run("Vraag").bold = True
            doc.add_paragraph(check)

        doc.add_paragraph("")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def docx_to_vmbo_lesson_json(file_like) -> io.BytesIO:
    """
    Hoofdfunctie die je in app.py importeert.
    - leest docx
    - stuurt ALLES naar AI
    - bouwt Word in les-format
    - GEEN fallback
    """
    blocks = docx_to_blocks(file_like)
    slides = ai_blocks_to_slides(blocks)
    return build_word_from_slides(slides)
